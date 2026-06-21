#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera Guía de usuario Chagas Tracker en Word y PDF en el Escritorio."""

import os
from datetime import date

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")
DOCX_PATH = os.path.join(DESKTOP, "Guia_Chagas_Tracker.docx")
PDF_PATH = os.path.join(DESKTOP, "Guia_Chagas_Tracker.pdf")

TITLE = "Chagas Tracker — Guía de uso del sistema"
SUBTITLE = (
    "Registro epidemiológico territorial anónimo de casos de Chagas\n"
    "Monte Patria, Región de Coquimbo, Chile"
)
VERSION = f"Versión 1.0 · {date.today().strftime('%d/%m/%Y')}"

SECTIONS = [
    (
        "1. Introducción",
        [
            "Chagas Tracker es una plataforma web responsiva para el registro epidemiológico "
            "territorial de casos de enfermedad de Chagas. No es un sistema clínico: no almacena "
            "nombre completo, RUT completo, teléfono, correo ni dirección exacta del paciente.",
            "El sistema está orientado al equipo del programa Chagas y operadores autorizados "
            "para registrar, consultar y dar seguimiento epidemiológico por sector territorial.",
        ],
    ),
    (
        "2. Requisitos",
        [
            "Navegador: Chrome, Edge o Firefox (versión reciente).",
            "Conexión a internet estable (los datos se guardan en Supabase).",
            "Cuenta de usuario creada por el administrador del programa.",
            "Dispositivo: computador, tablet o celular (interfaz adaptada a pantalla pequeña).",
        ],
    ),
    (
        "3. Acceso e inicio de sesión",
        [
            "Abra la URL de la aplicación en el navegador (proporcionada por el administrador).",
            "Ingrese correo y contraseña asignados.",
            "Pulse Ingresar.",
            "Si aparece «Correo o contraseña incorrectos», revise las credenciales.",
            "Si aparece «Sin conexión con el servidor», revise WiFi o datos móviles; "
            "no es necesariamente un error de contraseña.",
        ],
    ),
    (
        "4. Navegación principal",
        [
            "Tras iniciar sesión hay cuatro secciones:",
            "• Inicio: resumen con totales, accesos rápidos y actividad reciente.",
            "• Nuevo caso: formulario de registro epidemiológico.",
            "• Ver casos: listado, búsqueda y filtros.",
            "• Perfil: cuenta del operador, información del sistema y cierre de sesión.",
            "En pantallas anchas el menú está a la izquierda; en móvil, abajo.",
            "La franja naranja «Sin conexión» indica que no hay internet; puede verse "
            "información ya cargada, pero no se podrán guardar cambios hasta recuperar la red.",
        ],
    ),
    (
        "5. Pantalla Inicio (dashboard)",
        [
            "Accesos rápidos: registrar nuevo caso, ver casos, abrir el último caso registrado.",
            "Estado actual (KPIs): conteos de Caso nuevo, Reingreso, Tratado y Total. "
            "Al tocar una tarjeta se abre Ver casos con ese filtro.",
            "Sector con más casos: atajo al listado filtrado por el sector con mayor cantidad.",
            "Actividad reciente: últimos casos registrados; toque una fila para ver el detalle.",
            "Para actualizar: deslice hacia abajo (pull-to-refresh) o vuelva a la pestaña Inicio.",
        ],
    ),
    (
        "6. Registrar un nuevo caso",
        [
            "Complete solo datos epidemiológicos y territoriales.",
            "Sector (obligatorio): elija uno de los sectores del piloto:",
            "  · Chañaral Alto",
            "  · El Palqui",
            "  · Monte Patria",
            "  · Carén",
            "Identificador parcial: últimos 3 dígitos del documento + dígito verificador (ej. 123-K). "
            "No ingrese el RUT completo.",
            "Género y fecha de nacimiento: para caracterización epidemiológica.",
            "Ocupación: seleccione del catálogo; si no aplica, use «No informa».",
            "Número de contactos: cantidad epidemiológica (no nombres ni teléfonos).",
            "Observaciones: texto libre sin datos personales identificables.",
            "Duplicados: si existe un caso similar, el sistema ofrece Cancelar, Ver existente "
            "o Guardar de todos modos.",
            "Pulse Guardar caso; tras el éxito se abre la ficha del caso.",
        ],
    ),
    (
        "7. Ver casos (listado)",
        [
            "Búsqueda: código, sector, ocupación, rango etario o identificador parcial.",
            "Filtros rápidos por estado: Todos, Caso nuevo, Reingreso, Tratado.",
            "Chips de filtros activos arriba; use Limpiar para quitar todos.",
            "Filtros avanzados: género, sector, rango de fechas, orden por fecha.",
            "Toque una tarjeta para abrir el detalle del caso.",
            "Deslice hacia abajo para recargar la lista desde el servidor.",
        ],
    ),
    (
        "8. Detalle de un caso",
        [
            "Muestra identificación epidemiológica, ubicación territorial (sector/comuna), "
            "estado, historial de cambios y observaciones.",
            "Acciones en la barra inferior:",
            "  · Cambiar estado: Caso nuevo, Reingreso o Tratado.",
            "  · Editar datos: identificador parcial, género, fecha nac., ocupación, contactos.",
            "  · Editar observación: actualizar texto (sin datos personales).",
            "  · Exportar (icono): generar PDF del registro.",
            "Al volver al listado tras editar, la lista se actualiza automáticamente.",
        ],
    ),
    (
        "9. Perfil",
        [
            "Muestra correo y rol del operador autenticado.",
            "Acerca de Chagas Tracker: descripción breve del sistema.",
            "Buenas prácticas de registro: recordatorio de privacidad.",
            "Cerrar sesión: finaliza la sesión; use en equipos compartidos.",
        ],
    ),
    (
        "10. Buenas prácticas de privacidad",
        [
            "Use solo identificador parcial (3 dígitos + DV), nunca RUT completo.",
            "No registre nombres, teléfonos, correos ni direcciones en observaciones.",
            "El sector es información territorial agregada, no domicilio exacto.",
            "No comparta capturas con datos sensibles fuera del equipo autorizado.",
            "Cierre sesión al terminar en equipos compartidos.",
        ],
    ),
    (
        "11. Problemas frecuentes",
        [
            "Pantalla en blanco: espere la carga o recargue (F5).",
            "KPIs en cero: vaya a Inicio y espere recarga, o pull-to-refresh.",
            "Catálogo de ocupaciones vacío: revise conexión e intente de nuevo.",
            "Sesión expirada: vuelva a iniciar sesión.",
            "PDF no abre: pruebe otro navegador o permita ventanas emergentes.",
        ],
    ),
    (
        "12. Soporte",
        [
            "Consultas técnicas o altas de usuario: contactar al administrador del "
            "programa Chagas o al responsable del proyecto.",
        ],
    ),
]


def build_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading(TITLE, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(SUBTITLE.replace("\n", "\n"))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver = doc.add_paragraph(VERSION)
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ver.runs:
        ver.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for p in paragraphs:
            doc.add_paragraph(p, style="List Bullet" if p.startswith("•") or p.startswith("  ·") else "Normal")

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Chagas Tracker — Registro epidemiológico territorial. "
        "Documento generado para uso institucional."
    )
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if foot.runs:
        foot.runs[0].italic = True
        foot.runs[0].font.size = Pt(9)

    doc.save(DOCX_PATH)
    print(f"Word: {DOCX_PATH}")


def _windows_fonts_dir():
    windir = os.environ.get("WINDIR", r"C:\Windows")
    return os.path.join(windir, "Fonts")


def build_pdf():
    from fpdf import FPDF

    fonts = _windows_fonts_dir()
    arial = os.path.join(fonts, "arial.ttf")
    arial_b = os.path.join(fonts, "arialbd.ttf")
    arial_i = os.path.join(fonts, "ariali.ttf")
    if not os.path.isfile(arial):
        raise FileNotFoundError(
            f"No se encontró Arial en {fonts}. Instale fuentes del sistema o use el .docx."
        )

    class ManualPDF(FPDF):
        def footer(self):
            self.set_y(-15)
            self.set_font("Arial", "I", 8)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, f"Página {self.page_no()}", align="C")

    pdf = ManualPDF()
    pdf.add_font("Arial", "", arial)
    pdf.add_font("Arial", "B", arial_b)
    pdf.add_font("Arial", "I", arial_i)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    pdf.set_font("Arial", "B", 16)
    pdf.multi_cell(0, 9, TITLE, align="C")
    pdf.ln(4)
    pdf.set_font("Arial", "", 11)
    pdf.set_text_color(80, 80, 80)
    pdf.multi_cell(0, 6, SUBTITLE, align="C")
    pdf.ln(2)
    pdf.set_font("Arial", "I", 10)
    pdf.multi_cell(0, 6, VERSION, align="C")
    pdf.ln(8)
    pdf.set_text_color(0, 0, 0)

    for heading, paragraphs in SECTIONS:
        pdf.set_font("Arial", "B", 13)
        pdf.multi_cell(0, 8, heading)
        pdf.ln(2)
        pdf.set_font("Arial", "", 11)
        for p in paragraphs:
            pdf.multi_cell(0, 6, p)
            pdf.ln(1)
        pdf.ln(4)

    pdf.set_font("Arial", "I", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.multi_cell(
        0,
        5,
        "Chagas Tracker — Registro epidemiológico territorial. "
        "Documento generado para uso institucional.",
        align="C",
    )

    pdf.output(PDF_PATH)
    print(f"PDF: {PDF_PATH}")


if __name__ == "__main__":
    os.makedirs(DESKTOP, exist_ok=True)
    build_docx()
    build_pdf()
    print("Listo.")
